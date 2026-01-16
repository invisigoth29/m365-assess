#!/usr/bin/env python3
"""
Generate report_template.docx with docxtemplater tags
This creates a professional Word template for m365-assess reports
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_template():
    doc = Document()
    
    # Configure styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_font = title_style.font
    title_font.size = Pt(24)
    title_font.bold = True
    title_font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 1 style
    h1_style = styles['Heading 1']
    h1_font = h1_style.font
    h1_font.size = Pt(18)
    h1_font.color.rgb = RGBColor(0, 51, 102)
    
    # Heading 2 style
    h2_style = styles['Heading 2']
    h2_font = h2_style.font
    h2_font.size = Pt(14)
    h2_font.color.rgb = RGBColor(0, 102, 204)
    
    # ==== COVER PAGE ====
    title = doc.add_paragraph('Microsoft 365 Security Assessment', style='Title')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('{customer_name}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(68, 68, 68)
    
    date_para = doc.add_paragraph('{assessment_date}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para_run = date_para.runs[0]
    date_para_run.font.size = Pt(12)
    date_para_run.font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()  # Spacer
    
    prepared_by = doc.add_paragraph('Prepared by: {team_name}')
    prepared_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prepared_by_run = prepared_by.runs[0]
    prepared_by_run.font.size = Pt(11)
    prepared_by_run.font.italic = True
    
    doc.add_page_break()
    
    # ==== EXECUTIVE SUMMARY ====
    doc.add_heading('Executive Summary', 1)
    
    doc.add_paragraph(
        f'This report presents the findings of a comprehensive security assessment of '
        f'{{customer_name}}\'s Microsoft 365 environment conducted on {{assessment_date}}. '
        f'The assessment evaluated security configurations across identity management, email '
        f'protection, data security, and collaboration settings against Microsoft and CISA security baselines.'
    )
    
    doc.add_heading('Overall Security Posture', 2)
    
    # Security Score
    score_para = doc.add_paragraph()
    score_para.add_run('Security Score: ').bold = True
    score_para.add_run('{security_score}%')
    
    risk_para = doc.add_paragraph()
    risk_para.add_run('Risk Level: ').bold = True
    risk_para.add_run('{risk_level}')
    
    doc.add_paragraph()
    doc.add_paragraph('{assessment_summary}')
    
    # Findings Summary
    doc.add_heading('Findings Summary', 2)
    
    summary_para = doc.add_paragraph()
    summary_para.add_run(f'Total Controls Assessed: ').bold = True
    summary_para.add_run('{total_findings}')
    
    passed_para = doc.add_paragraph()
    passed_para.add_run('✓ Passed: ').bold = True
    passed_para.add_run('{passed_count}')
    
    failed_para = doc.add_paragraph()
    failed_para.add_run('✗ Failed: ').bold = True
    failed_para.add_run('{failed_count}')
    
    na_para = doc.add_paragraph()
    na_para.add_run('○ Not Applicable: ').bold = True
    na_para.add_run('{na_count}')
    
    # Risk breakdown
    doc.add_heading('Failed Findings by Severity', 3)
    
    doc.add_paragraph('Critical: {critical_count}')
    doc.add_paragraph('High: {high_count}')
    doc.add_paragraph('Medium: {medium_count}')
    doc.add_paragraph('Low: {low_count}')
    
    # Score movement (conditional)
    doc.add_heading('Score Movement', 2)
    doc.add_paragraph('{score_movement_message}')
    
    # Priority themes
    doc.add_heading('Priority Areas', 2)
    
    doc.add_paragraph(
        'Based on the assessment findings, the following security themes have been prioritized '
        'for remediation:'
    )
    
    # Loop through priority themes
    doc.add_paragraph('{#priority_themes}')
    theme_item = doc.add_paragraph('{title} ({priority}) - {failed_count} failed controls', style='List Bullet')
    doc.add_paragraph('{/priority_themes}')
    
    doc.add_page_break()
    
    # ==== REMEDIATION ROADMAP ====
    doc.add_heading('Remediation Roadmap', 1)
    
    doc.add_paragraph(
        'The following roadmap outlines a phased approach to addressing identified security gaps, '
        'organized by priority level and recommended timeline.'
    )
    
    # Roadmap table
    doc.add_heading('Prioritized Themes', 2)
    
    doc.add_paragraph('{#roadmap}')
    roadmap_item = doc.add_paragraph()
    roadmap_item.add_run('{priority}: ').bold = True
    roadmap_item.add_run('{theme} ({window}) - {failed_count} findings')
    doc.add_paragraph('{/roadmap}')
    
    doc.add_page_break()
    
    # ==== DETAILED FINDINGS BY THEME ====
    doc.add_heading('Detailed Findings by Security Theme', 1)
    
    doc.add_paragraph(
        'This section provides detailed analysis of each security theme, including business context, '
        'specific findings, and actionable remediation guidance.'
    )
    
    # Loop through themes
    doc.add_paragraph('{#themes}')
    
    # Theme header
    doc.add_heading('{title}', 2)
    
    # Theme metadata
    meta_para = doc.add_paragraph()
    meta_para.add_run('Priority: ').bold = True
    meta_para.add_run('{priority} ({window})')
    
    risk_para = doc.add_paragraph()
    risk_para.add_run('Risk Level: ').bold = True
    risk_para.add_run('{risk_level}')
    
    pass_rate_para = doc.add_paragraph()
    pass_rate_para.add_run('Pass Rate: ').bold = True
    pass_rate_para.add_run('{pass_rate}%')
    
    # Business context
    doc.add_heading('Business Rationale', 3)
    doc.add_paragraph('{business_rationale}')
    
    doc.add_heading('Business Impact', 3)
    doc.add_paragraph('{business_impact}')
    
    # Failed findings for this theme
    doc.add_heading('Failed Controls ({failed_findings.length})', 3)
    
    doc.add_paragraph('{#failed_findings}')
    finding_para = doc.add_paragraph('{control_id}: {title} (Severity: {severity})', style='List Bullet')
    doc.add_paragraph('{/failed_findings}')
    
    # Recommendations
    doc.add_heading('Recommendation Summary', 3)
    doc.add_paragraph('{recommendation_summary}')
    
    doc.add_heading('Remediation Steps', 3)
    doc.add_paragraph('{#remediation_steps}')
    doc.add_paragraph('{.}', style='List Number')
    doc.add_paragraph('{/remediation_steps}')
    
    doc.add_heading('Operational Notes', 3)
    doc.add_paragraph('{#operational_notes}')
    doc.add_paragraph('{.}', style='List Bullet')
    doc.add_paragraph('{/operational_notes}')
    
    doc.add_paragraph()  # Spacer between themes
    
    doc.add_paragraph('{/themes}')
    
    doc.add_page_break()
    
    # ==== APPENDIX ====
    doc.add_heading('Appendix', 1)
    
    doc.add_heading('Assessment Methodology', 2)
    doc.add_paragraph(
        'This assessment was conducted using ScubaGear, an automated security configuration '
        'assessment tool developed by CISA. ScubaGear evaluates Microsoft 365 configurations '
        'against the Secure Cloud Business Applications (SCuBA) security baseline.'
    )
    
    doc.add_heading('Assessment Metadata', 2)
    doc.add_paragraph(f'Tenant ID: {{tenant_id}}')
    doc.add_paragraph(f'Assessment Date: {{assessment_date}}')
    doc.add_paragraph(f'Run ID: {{run_id}}')
    doc.add_paragraph(f'Tool: ScubaGear via m365-assess wrapper')
    
    # Save template
    output_path = '/Users/donovanfarrell/Git/m365-assess/templates/report_template.docx'
    doc.save(output_path)
    print(f'✓ Template created: {output_path}')
    print('\nDocxtemplater tags used:')
    print('  - Simple variables: {customer_name}, {security_score}, etc.')
    print('  - Arrays: {#themes}...{/themes}, {#failed_findings}...{/failed_findings}')
    print('  - Array properties: {title}, {priority}, {severity}, etc.')
    print('  - Array length: {failed_findings.length}')
    print('  - Array items: {.} for simple array values')

if __name__ == '__main__':
    try:
        create_template()
    except Exception as e:
        print(f'Error: {e}')
        import traceback
        traceback.print_exc()
